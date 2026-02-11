"""Document processing service — PDF, XML (MedQuAD), and DOCX extraction.

Ported and improved from the original app.py:
  - extract_text_from_pdf (lines 116-128)  → now uses PyMuPDF (fitz)
  - extract_medquad_all  (lines 70-114)    → preserved logic
  - NEW: DOCX support via python-docx
  - NEW: OCR fallback via pytesseract for image-only PDF pages
"""
from __future__ import annotations

import logging
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ExtractedDocument:
    """Result of processing a single document."""
    source_filename: str
    file_type: str
    text: str = ""
    page_count: int = 0
    metadata: dict = field(default_factory=dict)
    error: str | None = None


@dataclass
class MedQuADPair:
    """A single Q&A pair extracted from a MedQuAD XML file."""
    question: str
    answer: str
    source_file: str
    quality_status: str = "Verified - Human Labeled"


# ── PDF Extraction ──────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path: str | Path) -> ExtractedDocument:
    """Extract text from a PDF using PyMuPDF (fitz), with OCR fallback.

    Falls back to pytesseract OCR if a page yields no text
    (i.e., scanned / image-only page).
    """
    pdf_path = Path(pdf_path)
    result = ExtractedDocument(source_filename=pdf_path.name, file_type="pdf")

    try:
        import fitz  # PyMuPDF
    except ImportError:
        # Fallback to PyPDF2 if fitz not installed
        return _extract_text_from_pdf_pypdf2(pdf_path, result)

    try:
        doc = fitz.open(str(pdf_path))
        result.page_count = len(doc)
        pages_text: list[str] = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text and text.strip():
                pages_text.append(text)
            else:
                # Try OCR fallback for image-only pages
                ocr_text = _ocr_page(page, page_num, pdf_path.name)
                if ocr_text:
                    pages_text.append(ocr_text)

        doc.close()
        result.text = "\n\n".join(pages_text)
        result.metadata = {"page_count": result.page_count, "method": "pymupdf"}

    except Exception as e:
        logger.error(f"Error processing PDF {pdf_path}: {e}")
        result.error = str(e)

    return result


def _extract_text_from_pdf_pypdf2(pdf_path: Path, result: ExtractedDocument) -> ExtractedDocument:
    """Fallback PDF extraction using PyPDF2 (legacy compat)."""
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            result.page_count = len(reader.pages)
            pages_text = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                    pages_text.append(text)
                except Exception as e:
                    logger.warning(f"PyPDF2 error on page {page_num} of {pdf_path}: {e}")
            result.text = "\n\n".join(pages_text)
            result.metadata = {"page_count": result.page_count, "method": "pypdf2"}
    except Exception as e:
        logger.error(f"PyPDF2 error on {pdf_path}: {e}")
        result.error = str(e)
    return result


def _ocr_page(page, page_num: int, filename: str) -> str:
    """Attempt OCR on a PyMuPDF page using pytesseract."""
    try:
        import pytesseract
        from PIL import Image
        import io

        # Render page to image at 300 DPI
        pix = page.get_pixmap(dpi=300)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        text = pytesseract.image_to_string(img)
        if text and text.strip():
            logger.info(f"OCR extracted text from page {page_num} of {filename}")
            return text.strip()
    except ImportError:
        logger.debug("pytesseract not available, skipping OCR")
    except Exception as e:
        logger.warning(f"OCR failed on page {page_num} of {filename}: {e}")
    return ""


# ── XML / MedQuAD Extraction ───────────────────────────────────────────

def extract_medquad_pairs(
    xml_path: str | Path,
    keywords: list[str],
) -> list[MedQuADPair]:
    """Extract Q&A pairs from a MedQuAD XML file filtered by keywords.

    Ported from original ``extract_medquad_all()`` — preserves the same
    keyword-matching logic that checks title, question, and answer text.
    """
    xml_path = Path(xml_path)
    pairs: list[MedQuADPair] = []
    lower_keywords = [kw.lower().strip() for kw in keywords if kw.strip()]

    if not lower_keywords:
        return pairs

    try:
        tree = ET.parse(str(xml_path))
        root = tree.getroot()

        for doc in root.findall(".//document"):
            title_elem = doc.find("full_title")
            title = title_elem.text if title_elem is not None and title_elem.text else ""

            qa_sections = doc.findall(".//qa_pairs/qa_pair") or doc.findall(".//qa_pair")
            for pair_elem in qa_sections:
                q_elem = pair_elem.find("question")
                a_elem = pair_elem.find("answer")
                question = q_elem.text.strip() if q_elem is not None and q_elem.text else ""
                answer = a_elem.text.strip() if a_elem is not None and a_elem.text else ""

                if not question or not answer:
                    continue

                combined = f"{title} {question} {answer}".lower()
                if any(kw in combined for kw in lower_keywords):
                    pairs.append(MedQuADPair(
                        question=question,
                        answer=answer,
                        source_file=xml_path.name,
                    ))

    except Exception as e:
        logger.error(f"Error parsing XML {xml_path}: {e}")

    return pairs


def extract_text_from_xml(xml_path: str | Path) -> ExtractedDocument:
    """Extract raw text content from an XML file (non-MedQuAD, generic)."""
    xml_path = Path(xml_path)
    result = ExtractedDocument(source_filename=xml_path.name, file_type="xml")
    try:
        tree = ET.parse(str(xml_path))
        root = tree.getroot()
        texts = [elem.text for elem in root.iter() if elem.text and elem.text.strip()]
        result.text = "\n\n".join(texts)
    except Exception as e:
        logger.error(f"Error extracting text from XML {xml_path}: {e}")
        result.error = str(e)
    return result


# ── DOCX Extraction ────────────────────────────────────────────────────

def extract_text_from_docx(docx_path: str | Path) -> ExtractedDocument:
    """Extract text from a Word document using python-docx."""
    docx_path = Path(docx_path)
    result = ExtractedDocument(source_filename=docx_path.name, file_type="docx")
    try:
        import docx
        doc = docx.Document(str(docx_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result.text = "\n\n".join(paragraphs)
        result.metadata = {"paragraph_count": len(paragraphs)}
    except ImportError:
        result.error = "python-docx not installed"
        logger.warning("python-docx not installed; cannot process DOCX files")
    except Exception as e:
        logger.error(f"Error processing DOCX {docx_path}: {e}")
        result.error = str(e)
    return result


# ── Dispatcher ─────────────────────────────────────────────────────────

def extract_text(filepath: str | Path) -> ExtractedDocument:
    """Dispatch to the appropriate extractor based on file extension."""
    filepath = Path(filepath)
    ext = filepath.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".xml":
        return extract_text_from_xml(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        return ExtractedDocument(
            source_filename=filepath.name,
            file_type=ext.lstrip("."),
            error=f"Unsupported file type: {ext}",
        )
